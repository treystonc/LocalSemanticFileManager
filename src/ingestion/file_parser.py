"""File parser module for extracting text from various file formats.

Supports: .pdf, .docx, .xlsx, .txt, .md
"""

import logging
from pathlib import Path
from typing import Optional

import pymupdf
import pandas as pd
from docx import Document

logger = logging.getLogger(__name__)


class FileParser:
    """Extracts text content from various file formats."""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".md"}
    
    def __init__(self, supported_extensions: Optional[list[str]] = None) -> None:
        """Initialize the file parser.
        
        Args:
            supported_extensions: List of extensions to support (default: all)
        """
        if supported_extensions:
            self.supported_extensions = set(
                ext.lower() if ext.startswith(".") else f".{ext.lower()}"
                for ext in supported_extensions
            )
        else:
            self.supported_extensions = self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: Path) -> dict:
        """Parse a file and extract its text content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with 'text', 'metadata', and 'success' keys
        """
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        
        result = {
            "text": "",
            "metadata": {
                "filename": file_path.name,
                "path": str(file_path),
                "extension": ext,
                "size_bytes": file_path.stat().st_size if file_path.exists() else 0,
            },
            "success": False,
            "error": None,
        }
        
        if not file_path.exists():
            result["error"] = "File does not exist"
            return result
        
        if ext not in self.supported_extensions:
            result["error"] = f"Unsupported extension: {ext}"
            return result
        
        parser_map = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".xlsx": self._parse_xlsx,
            ".txt": self._parse_text,
            ".md": self._parse_text,
        }
        
        parser_func = parser_map.get(ext)
        if not parser_func:
            result["error"] = f"No parser for extension: {ext}"
            return result
        
        try:
            text = parser_func(file_path)
            result["text"] = text
            result["success"] = True
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            result["error"] = str(e)
        
        return result
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF."""
        text_parts = []
        
        with pymupdf.open(str(file_path)) as doc:
            for page in doc:
                text = page.get_text()
                if text.strip():
                    text_parts.append(text.strip())
        
        return "\n\n".join(text_parts)
    
    def _parse_docx(self, file_path: Path) -> str:
        """Extract text from Word document."""
        doc = Document(str(file_path))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    def _parse_xlsx(self, file_path: Path) -> str:
        """Extract text from Excel file."""
        text_parts = []
        
        xlsx = pd.ExcelFile(str(file_path))
        
        for sheet_name in xlsx.sheet_names:
            df = pd.read_excel(xlsx, sheet_name=sheet_name, header=None)
            text_parts.append(f"[Sheet: {sheet_name}]")
            
            for _, row in df.iterrows():
                row_text = " | ".join(
                    str(cell) for cell in row 
                    if pd.notna(cell) and str(cell).strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    def _parse_text(self, file_path: Path) -> str:
        """Extract text from plain text files."""
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode file with any supported encoding")
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if a file is supported."""
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def parse_directory(
        self, 
        directory: Path, 
        recursive: bool = False
    ) -> list[dict]:
        """Parse all supported files in a directory.
        
        Args:
            directory: Directory to scan
            recursive: Whether to scan subdirectories
            
        Returns:
            List of parse results
        """
        directory = Path(directory)
        results = []
        
        if not directory.exists():
            logger.error(f"Directory does not exist: {directory}")
            return results
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory.glob(pattern):
            if file_path.is_file() and self.is_supported(file_path):
                result = self.parse(file_path)
                results.append(result)
        
        return results
